"""
Document AI Module.

Provides document text extraction, summarization, and deadline extraction.
Supports PDF, Word documents, and plain text.
"""

import os
import re
import logging
import time
from typing import Optional, Dict, Any, List, Tuple
from pathlib import Path
from datetime import datetime

logger = logging.getLogger(__name__)


def extract_text_from_file(file_path: str) -> Tuple[str, str]:
    """
    Extract text content from a document file.

    Args:
        file_path: Path to the document file

    Returns:
        Tuple of (extracted_text, file_type)

    Raises:
        ValueError: If file type is not supported
        FileNotFoundError: If file doesn't exist
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = path.suffix.lower()

    if suffix == '.pdf':
        return extract_pdf_text(file_path), 'pdf'
    elif suffix in ['.docx', '.doc']:
        return extract_docx_text(file_path), 'docx'
    elif suffix == '.txt':
        return extract_txt_text(file_path), 'txt'
    elif suffix in ['.md', '.markdown']:
        return extract_txt_text(file_path), 'markdown'
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def extract_pdf_text(file_path: str) -> str:
    """
    Extract text from a PDF file.

    Uses PyPDF2 for text extraction. Falls back to empty string
    if extraction fails (e.g., scanned PDFs).
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            logger.error("Neither pypdf nor PyPDF2 is installed")
            return ""

    try:
        reader = PdfReader(file_path)
        text_parts = []

        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        return '\n\n'.join(text_parts)

    except Exception as e:
        logger.error(f"Failed to extract PDF text: {e}")
        return ""


def extract_docx_text(file_path: str) -> str:
    """
    Extract text from a Word document (.docx).
    """
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx is not installed")
        return ""

    try:
        doc = Document(file_path)
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return '\n\n'.join(text_parts)

    except Exception as e:
        logger.error(f"Failed to extract DOCX text: {e}")
        return ""


def extract_txt_text(file_path: str) -> str:
    """
    Extract text from a plain text file.
    """
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception as e:
        logger.error(f"Failed to read text file: {e}")
        return ""


async def summarize_document(
    text: str,
    max_chars: int = 50000,
    db=None,
    organization_id: int = None
) -> Optional[Dict[str, Any]]:
    """
    Generate an AI summary of document text.

    Args:
        text: Document text content
        max_chars: Maximum characters to process
        db: Database session for provider preferences
        organization_id: Organization ID for preferences

    Returns:
        Dict with summary, key_terms, dates, risk_flags, etc.
    """
    from .llm_client import parse_json_response
    from .providers import get_fallback_client
    from .prompts import DOCUMENT_PROMPTS

    if not text or len(text.strip()) < 50:
        return None

    # Truncate if too long
    if len(text) > max_chars:
        text = text[:max_chars] + "\n\n[Document truncated for summarization]"

    # Use fallback client with preference for local (cost savings)
    client = get_fallback_client(db, organization_id, prefer_local=True)
    prompt = DOCUMENT_PROMPTS["summarize"].format(content=text)

    start_time = time.time()
    response = await client.generate(
        prompt=prompt,
        temperature=0.2,
        max_tokens=2048
    )
    processing_time = int((time.time() - start_time) * 1000)

    if not response.success:
        logger.error(f"Document summarization failed: {response.error}")
        return None

    parsed = parse_json_response(response.content)

    if parsed:
        return {
            "summary": parsed.get("summary", ""),
            "document_type": parsed.get("document_type", "other"),
            "key_terms": parsed.get("key_terms", []),
            "extracted_dates": parsed.get("dates", []),
            "action_items": parsed.get("action_items", []),
            "risk_flags": parsed.get("risk_flags", []),
            "tokens_used": response.tokens_used,
            "model_used": response.model,
            "processing_time_ms": processing_time
        }

    return None


async def extract_deadlines_from_text(
    text: str,
    max_chars: int = 30000,
    db=None,
    organization_id: int = None
) -> Optional[Dict[str, Any]]:
    """
    Extract dates and deadlines from document text.

    Args:
        text: Document text content
        max_chars: Maximum characters to process
        db: Database session for provider preferences
        organization_id: Organization ID for preferences

    Returns:
        Dict with deadlines and recurring_dates
    """
    from .llm_client import parse_json_response
    from .providers import get_fallback_client
    from .prompts import DOCUMENT_PROMPTS

    if not text or len(text.strip()) < 20:
        return None

    # Truncate if too long
    if len(text) > max_chars:
        text = text[:max_chars] + "\n\n[Document truncated]"

    # Use fallback client with preference for local (cost savings)
    client = get_fallback_client(db, organization_id, prefer_local=True)
    prompt = DOCUMENT_PROMPTS["extract_deadlines"].format(content=text)

    response = await client.generate(
        prompt=prompt,
        temperature=0.1,  # Lower temperature for more deterministic date extraction
        max_tokens=1024
    )

    if not response.success:
        logger.error(f"Deadline extraction failed: {response.error}")
        return None

    parsed = parse_json_response(response.content)

    if parsed:
        return {
            "deadlines": parsed.get("deadlines", []),
            "recurring_dates": parsed.get("recurring_dates", [])
        }

    return None


def extract_dates_regex(text: str) -> List[Dict[str, str]]:
    """
    Extract dates using regex patterns as a fallback.

    This is faster than LLM and useful for simple date detection.
    """
    dates = []

    # Common date patterns
    patterns = [
        # MM/DD/YYYY or MM-DD-YYYY
        (r'\b(\d{1,2})[/-](\d{1,2})[/-](\d{4})\b', 'us'),
        # YYYY-MM-DD (ISO)
        (r'\b(\d{4})-(\d{2})-(\d{2})\b', 'iso'),
        # Month DD, YYYY
        (r'\b(January|February|March|April|May|June|July|August|September|October|November|December)\s+(\d{1,2}),?\s+(\d{4})\b', 'full'),
        # DD Month YYYY
        (r'\b(\d{1,2})\s+(January|February|March|April|May|June|July|August|September|October|November|December)\s+(\d{4})\b', 'full_eu'),
    ]

    month_map = {
        'january': 1, 'february': 2, 'march': 3, 'april': 4,
        'may': 5, 'june': 6, 'july': 7, 'august': 8,
        'september': 9, 'october': 10, 'november': 11, 'december': 12
    }

    for pattern, fmt in patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            try:
                if fmt == 'us':
                    month, day, year = int(match.group(1)), int(match.group(2)), int(match.group(3))
                elif fmt == 'iso':
                    year, month, day = int(match.group(1)), int(match.group(2)), int(match.group(3))
                elif fmt == 'full':
                    month = month_map[match.group(1).lower()]
                    day = int(match.group(2))
                    year = int(match.group(3))
                elif fmt == 'full_eu':
                    day = int(match.group(1))
                    month = month_map[match.group(2).lower()]
                    year = int(match.group(3))

                # Validate date
                date_obj = datetime(year, month, day)
                date_str = date_obj.strftime('%Y-%m-%d')

                # Get surrounding context
                start = max(0, match.start() - 50)
                end = min(len(text), match.end() + 50)
                context = text[start:end].strip()

                dates.append({
                    "date": date_str,
                    "source_text": context,
                    "confidence": "high" if fmt in ['iso', 'full'] else "medium"
                })

            except (ValueError, KeyError):
                continue

    return dates


def detect_document_type(text: str) -> str:
    """
    Detect document type based on content analysis.

    Returns one of: contract, term_sheet, invoice, report, legal, letter, other
    """
    text_lower = text.lower()

    # Term sheet indicators
    if any(term in text_lower for term in ['term sheet', 'valuation cap', 'pre-money', 'post-money', 'liquidation preference']):
        return 'term_sheet'

    # Contract indicators
    if any(term in text_lower for term in ['agreement', 'hereby agrees', 'parties agree', 'effective date', 'termination']):
        return 'contract'

    # Invoice indicators
    if any(term in text_lower for term in ['invoice', 'bill to', 'due date', 'payment terms', 'total amount']):
        return 'invoice'

    # Legal document indicators
    if any(term in text_lower for term in ['whereas', 'hereinafter', 'pursuant to', 'notwithstanding']):
        return 'legal'

    # Report indicators
    if any(term in text_lower for term in ['executive summary', 'findings', 'analysis', 'recommendations', 'conclusions']):
        return 'report'

    # Letter indicators
    if any(term in text_lower for term in ['dear ', 'sincerely', 'best regards', 'yours truly']):
        return 'letter'

    return 'other'


def chunk_text(text: str, chunk_size: int = 4000, overlap: int = 200) -> List[str]:
    """
    Split text into chunks for processing long documents.

    Args:
        text: Full document text
        chunk_size: Maximum characters per chunk
        overlap: Overlap between chunks for context continuity

    Returns:
        List of text chunks
    """
    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0

    while start < len(text):
        end = start + chunk_size

        # Try to break at a paragraph or sentence boundary
        if end < len(text):
            # Look for paragraph break
            para_break = text.rfind('\n\n', start, end)
            if para_break > start + chunk_size // 2:
                end = para_break

            # Or sentence break
            elif text.rfind('. ', start, end) > start + chunk_size // 2:
                end = text.rfind('. ', start, end) + 1

        chunks.append(text[start:end].strip())
        start = end - overlap

    return chunks
